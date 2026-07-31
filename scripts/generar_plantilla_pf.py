#!/usr/bin/env python3
"""
generar_plantilla_pf.py — Build PF-G2-...docx and BEPF-G2-...docx with the
exact format required by the assignment (Arial 12, doble espacio, justificado,
una sola columna, sin hoja de presentación).

Sections 1-7 come from ANTE-PROYECTO.md (with verbs in past tense where
appropriate, per the instructions: "Desde Título hasta Justificación, estas
secciones son en principio las mismas que en el anteproyecto, salvo por algunos
cambios menores que se deban realizar para complementar lo hecho en el
anteproyecto.").

Sections 8-10 (Resultados, Mejoras, Conclusiones) are placeholders the team
will fill in.

Section 11 (Referencias APA) reproduces the 6 references from the anteproyecto.

Output:
  templates/PF-G2-Bello Ballen, Cristancho Niño, Gordillo Meneses, Vargas Clavijo.docx
  templates/BEPF-G2-Bello Ballen, Cristancho Niño, Gordillo Meneses, Vargas Clavijo.docx
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
GRAFICAS_DIR = PROJECT_ROOT / "graficas"
CAL_JSON = PROJECT_ROOT / "cal.json"
RMS_CSV = PROJECT_ROOT / "rms.csv"

TITLE = (
    "MODELO DE CRECIMIENTO LOGÍSTICO CON EXTRACCIÓN APLICADO A LA "
    "DINÁMICA DE REGENERACIÓN DE CÉSPED EN MINECRAFT"
)

AUTHORS = (
    "Bello Ballen Lina Andrea, Cristancho Niño Julián David, "
    "Gordillo Meneses Mariana Alejandra, Vargas Clavijo Brian Steven"
)

INTRO = (
    "Este proyecto aplica el modelo de crecimiento logístico con extracción "
    "para analizar la dinámica de regeneración de césped frente al consumo "
    "de ovejas en Minecraft, un entorno virtual con reglas algorítmicas "
    "fijas: el motor procesa 20 ticks por segundo (Minecraft Wiki, s.f.), "
    "las ovejas consumen bloques de césped con probabilidad constante "
    "(Sportskeeda, 2024) y el césped se regenera por propagación a bloques "
    "de tierra cercanos bajo condiciones de luz adecuadas (Minecraft Wiki, "
    "2024)."
)

GLOSSARY_PREFIX = "Notación: "
GLOSSARY_TEXT = (
    "G(t) = bloques de césped en el corral; r = tasa de regeneración (1/s); "
    "K = capacidad máxima del corral (bloques); c = tasa de consumo por oveja "
    "(bloques/oveja·s); N = número de ovejas; G*₊ = equilibrio estable "
    "(bloques); N_crit = ovejas en el punto de bifurcación silla-nodo."
)

# Estado del arte ahora se construye con bloques: cada bloque es un párrafo
# de texto o una ecuación. Esto permite intercalar las ecuaciones (1) y (2)
# centradas y numeradas, como pidió el docente.
ESTADO_BLOCKS = [
    ("text",
     "Para modelar mediante EDOs la rapidez con la que las ovejas "
     "consumen el césped dentro de un corral en Minecraft, considerando "
     "regeneración y consumo proporcional al número de ovejas, se "
     "seleccionaron como base teórica los artículos Harvesting Policies "
     "with Stepwise Effort and Logistic Growth in a Random Environment "
     "(Brites & Braumann, 2020) y Trimming to Coexistence (Braverman & "
     "Lawson, 2025), que abordan dinámica de poblaciones con EDOs "
     "adaptables al problema planteado."),
    ("text",
     "Brites y Braumann (2020) presentan un modelo logístico con "
     "extracción para poblaciones que crecen con límite y son "
     "simultáneamente reducidas por consumo; su forma general es:"),
    ("eq", r"\frac{dX}{dt} = rX\left(1 - \frac{X}{K}\right) - H(t)", 1),
    ("text",
     "donde X(t) es la población, r la tasa de crecimiento, K la "
     "capacidad del entorno y H(t) la tasa de extracción. En el sistema "
     "de Minecraft el césped corresponde al recurso con regeneración "
     "limitada y las ovejas representan el proceso de consumo, con "
     "extracción proporcional al número de ovejas. Por su parte, "
     "Braverman y Lawson (2025) estudian sistemas dinámicos "
     "recurso-consumidor mediante EDOs y aportan herramientas de "
     "análisis de estabilidad y equilibrio."),
    ("text",
     "Aplicando estas ideas, con G(t) la cantidad de césped disponible "
     "y N el número de ovejas, se obtuvo el modelo:"),
    ("eq", r"\frac{dG}{dt} = rG\left(1 - \frac{G}{K}\right) - cN", 2),
    ("text",
     "donde r es la tasa de regeneración, K el máximo de bloques con "
     "césped y c la tasa de consumo por oveja. La ecuación (2) permite "
     "analizar escenarios variando ovejas, tamaño del corral y "
     "obstáculos, y será objeto de calibración experimental y análisis "
     "sistemático en las siguientes secciones."),
]

PROPUESTA = (
    "El proyecto modeló el ecosistema de un corral en Minecraft como un "
    "sistema de recurso renovable sometido a consumo constante. Aunque los "
    "modelos logísticos con extracción se utilizan comúnmente para describir "
    "poblaciones biológicas o recursos naturales, su aplicación a sistemas "
    "virtuales gobernados por reglas algorítmicas ha sido menos explorada. En "
    "este contexto, el césped se interpretó como una población que se "
    "regenera con crecimiento logístico, mientras que las ovejas actuaron "
    "como agentes que consumen el recurso. El sistema se describió mediante "
    "una ecuación diferencial que incorpora la regeneración limitada por la "
    "capacidad del entorno y el consumo proporcional al número de ovejas. Con "
    "base en la teoría de modelos logísticos con cosecha (Alharbi, 2020), se "
    "analizó si el sistema alcanza un equilibrio sostenible o si el césped se "
    "agota cuando la extracción supera la capacidad de regeneración."
)

OBJETIVO_GENERAL = (
    "Desarrollar un modelo matemático basado en EDOs que permita analizar la "
    "dinámica entre el consumo y la regeneración del césped dentro de un "
    "corral en Minecraft, con el fin de determinar qué tan rápido se agota el "
    "recurso bajo distintos escenarios, variando el número de ovejas, el "
    "tamaño y el tipo de obstáculos del área cercada, evaluando así las "
    "condiciones en las que el sistema alcanza equilibrio o en las que el "
    "césped desaparece."
)

OBJETIVOS_ESPEC = [
    # 1 — consolidado: formulación + calibración (atiende comentario 4)
    "Formular un modelo basado en una EDO que describa la regeneración del "
    "césped y su consumo por las ovejas, y determinar experimentalmente los "
    "parámetros de regeneración r y consumo c mediante un protocolo de "
    "calibración controlada en un corral de referencia.",
    # 2 — consolida los antiguos 2, 3, 4, 5
    "Analizar el comportamiento del sistema bajo distintos escenarios — "
    "variando el número de ovejas, el tamaño del corral y el tipo de "
    "obstáculo presente (agua, hojas, vallas) — para identificar las "
    "condiciones bajo las cuales el sistema alcanza equilibrio sostenible y "
    "aquellas en las que el césped se agota.",
]

# Metodología también con bloques para intercalar la ecuación (3).
METODOLOGIA_BLOCKS = [
    ("step",
     "Sistema de estudio. Se diseñaron 12 configuraciones experimentales "
     "distribuidas en 6 corrales construidos en Minecraft: 3 regulares "
     "(5×5, 10×10 y 15×15) y 3 con obstáculos en una base 12×12. En estos "
     "últimos se incluyeron 44 bloques de obstáculo (agua, hojas o vallas) "
     "diseñados para preservar un área efectiva de césped de 100 bloques, "
     "comparable a la del corral regular C2 (10×10). Esto permitió aislar "
     "el efecto del tipo de obstáculo manteniendo K constante."),
    ("step",
     "Formulación del modelo. El sistema se parametrizó con variables "
     "continuas y se modeló mediante la EDO logística con extracción, "
     "ecuación (2)."),
    ("step",
     "Condiciones iniciales. Cada experimento inició con el terreno "
     "completamente cubierto de césped (G₀ = K) y con las ovejas "
     "distribuidas uniformemente dentro del corral."),
    ("step",
     "Calibración de parámetros. La tasa de regeneración r y la tasa de "
     "consumo c se determinaron experimentalmente en el corral C2 (K=100) "
     "mediante cinco réplicas independientes. Para r se observó la "
     "saturación G(t) → K partiendo de G₀ = 1 sin ovejas y se ajustó la "
     "solución analítica de la logística por mínimos cuadrados. Para c se "
     "ajustó la EDO (2) con N = 5 ovejas, manteniendo r fijo. Los valores "
     "obtenidos se asumieron constantes para los 12 escenarios."),
    ("step",
     "Captura automatizada de datos. La trayectoria G(t) se midió cada "
     "100 ticks de juego mediante command blocks que cuentan los bloques "
     "de césped en el corral activo. El registro queda almacenado en el "
     "archivo de log del servidor Minecraft, alojado en un VPS remoto. "
     "Una rutina automatizada conectada por SSH extrae los logs, los "
     "renombra según escenario y réplica, los procesa con un parser y "
     "produce un CSV unificado. Cada corrida se limitó a un máximo de "
     "150 000 ticks de juego (≈ 125 minutos al ritmo del servidor) tras "
     "lo cual se detuvo manualmente; corridas detenidas en el tope se "
     "siguen analizando sobre su curva parcial. Adicionalmente, se "
     "realizaron grabaciones de pantalla de algunas sesiones como "
     "respaldo visual de los eventos clave."),
    ("step",
     "Escenarios. Se variaron N (de 1 a 40) y el tipo de obstáculo "
     "manteniendo el área efectiva. Cada escenario se replicó hasta tres "
     "veces para promediar la estocasticidad de los random ticks."),
    ("step",
     "Análisis. La trayectoria experimental promedio se comparó con la "
     "solución numérica de la EDO y con el equilibrio analítico, ecuación "
     "(3), cuando el discriminante es no negativo. El error se cuantificó "
     "con la raíz cuadrática media (RMS) de la diferencia entre datos y "
     "modelo."),
    ("eq", r"G^*_+ = \frac{K}{2}\left(1 + \sqrt{1 - \frac{4cN}{rK}}\right)", 3),
    ("step",
     "Comparación. Finalmente, los escenarios se compararon para "
     "identificar bajo qué condiciones el sistema sostiene equilibrio o "
     "colapsa, y para aislar el efecto del obstáculo comparando los "
     "escenarios C2/C4/C5/C6 con K = 100 y N = 10 idénticos."),
]

JUSTIFICACION_PARRAFOS = [
    "Este proyecto surgió del interés por aplicar EDOs a un sistema "
    "virtual con reglas precisas y medibles. Minecraft ofrece un "
    "entorno controlado y reproducible, ideal para observar, registrar "
    "y modelar el comportamiento de un recurso renovable bajo consumo. "
    "Se tradujo un fenómeno del juego al lenguaje matemático del "
    "modelo logístico con extracción, permitiendo predecir cuándo el "
    "sistema es sostenible y cuándo colapsa, lo que hace visible la "
    "utilidad de las EDOs en un contexto tangible. Adicionalmente, el "
    "proyecto integró infraestructura computacional moderna: la "
    "simulación corrió en un servidor Minecraft alojado en un VPS y la "
    "captura de datos se automatizó mediante una rutina basada en "
    "conexiones SSH, extracción de logs y análisis numérico, "
    "demostrando que una herramienta matemática se puede convertir en "
    "un pipeline reproducible.",
]

# Referencias: cada entry es lista de runs (text, italic_bool) para que el
# nombre del journal y el volumen queden en cursiva (norma APA 7).
REFERENCIAS = [
    [("Alharbi, F. M. (2020). Analysis of logistic model with constant "
      "harvesting in a view of non-integer derivative. ", False),
     ("Journal of Mathematics and System Science, 10", True),
     ("(2). https://doi.org/10.17265/2159-5291/2020.02.004", False)],
    [("Braverman, E., & Lawson, J. (2025). Trimming to coexistence: How "
      "dispersal strategies should be accounted for in resource management. ",
      False),
     ("Journal of Mathematical Biology, 92", True),
     ("(1), 8. https://doi.org/10.1007/s00285-025-02324-8", False)],
    [("Brites, N. M., & Braumann, C. A. (2020). Harvesting policies with "
      "stepwise effort and logistic growth in a random environment. ", False),
     ("Applied Stochastic Models in Business and Industry, 36", True),
     ("(5), 825–835. https://doi.org/10.1002/asmb.2532", False)],
    [("Holling, C. S. (1959). The components of predation as revealed by a "
      "study of small-mammal predation of the European pine sawfly. ", False),
     ("The Canadian Entomologist, 91", True),
     ("(5), 293–320. https://doi.org/10.4039/Ent91293-5", False)],
    [("Minecraft Wiki. (2024, 25 de octubre). Grass block. "
      "https://minecraft.wiki/w/Grass_Block", False)],
    [("Minecraft Wiki. (s. f.). Tick. https://es.minecraft.wiki/w/Tick", False)],
    [("Real, L. A. (1977). The kinetics of functional response. ", False),
     ("The American Naturalist, 111", True),
     ("(978), 289–300. https://doi.org/10.1086/283161", False)],
    [("Sportskeeda. (2024, 24 de mayo). Sheep Minecraft. "
      "https://wiki.sportskeeda.com/minecraft/sheep", False)],
    [("Turchin, P. (2003). ", False),
     ("Complex population dynamics: A theoretical/empirical synthesis", True),
     (". Princeton University Press.", False)],
]


def setup_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:cs"), "Arial")
    rfonts.set(qn("w:eastAsia"), "Arial")
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            from docx.oxml import OxmlElement
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Arial")
        rfonts.set(qn("w:hAnsi"), "Arial")
        rfonts.set(qn("w:cs"), "Arial")
    h.paragraph_format.line_spacing = 2.0


def add_placeholder(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0


def add_equation(doc: Document, formula: str, num: int) -> None:
    """Insert a centered equation paragraph with right-aligned number `(num)`.
    Fallback de texto plano (mantenido por compatibilidad)."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(15.24), WD_TAB_ALIGNMENT.RIGHT
    )
    p.add_run(formula)
    p.add_run(f"\t({num})")


def _render_latex_to_png(latex_src: str, fontsize: int = 16,
                          dpi: int = 200) -> bytes:
    """Renderizar ecuación LaTeX a PNG via matplotlib mathtext.

    Genera ecuaciones con tipografía matemática profesional (similar a
    LaTeX puro) usando el motor mathtext de matplotlib. Salida transparente.
    """
    import io
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig = plt.figure(figsize=(0.01, 0.01))
    fig.text(0, 0, f"${latex_src}$", fontsize=fontsize)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                pad_inches=0.05, transparent=True)
    plt.close(fig)
    buf.seek(0)
    return buf


def add_equation_latex(doc: Document, latex_src: str, num: int,
                        height_pt: float = 22.0) -> None:
    """Inserta ecuación renderizada vía LaTeX (matplotlib mathtext) +
    numeración (N) a la derecha. Norma APA: ecuación centrada con número
    de identificación entre paréntesis al margen derecho.
    """
    png = _render_latex_to_png(latex_src)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(15.24), WD_TAB_ALIGNMENT.RIGHT
    )
    run_img = p.add_run()
    run_img.add_picture(png, height=Pt(height_pt))
    run_num = p.add_run(f"\t({num})")
    run_num.font.size = Pt(11)


def add_glossary(doc: Document) -> None:
    """Pre-poblar el glosario de variables como párrafo justificado al final
    de la Introducción. Prefix 'Notación:' en negrita. Mariana puede
    convertirlo a footnote real en Word si quiere — python-docx no soporta
    footnotes nativamente."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    bold_run = p.add_run(GLOSSARY_PREFIX)
    bold_run.bold = True
    p.add_run(GLOSSARY_TEXT)


def add_picture_with_caption(doc: Document, img_path: Path, caption: str,
                              width_inches: float = 6.0) -> None:
    """Insertar imagen centrada + leyenda 'Figura N — ...' en cursiva."""
    if not img_path.exists():
        # fallback: marcador de texto
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[{img_path.name} no encontrada]")
        run.italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.add_run().add_picture(str(img_path), width=Inches(width_inches))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.15
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)


def add_table_from_rows(doc: Document, headers: list[str],
                         rows: list[list[str]], caption: str) -> None:
    """Tabla simple con encabezado + filas + leyenda 'Tabla N — ...'."""
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.15
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # encabezado
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        para = hdr_cells[i].paragraphs[0]
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    # filas
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(str(val))
            run.font.size = Pt(10)
    # espacio post-tabla
    doc.add_paragraph()


def load_cal_and_rms() -> tuple[dict, pd.DataFrame]:
    """Lee cal.json y rms.csv del proyecto. Si no existen, retorna defaults."""
    if CAL_JSON.exists():
        cal = json.loads(CAL_JSON.read_text())
    else:
        cal = {"r_cal": 0.005037, "c_cal": 0.010355}
    if RMS_CSV.exists():
        rms = pd.read_csv(RMS_CSV)
    else:
        rms = pd.DataFrame()
    return cal, rms


def _render_resultados(doc: Document) -> None:
    """Resultados completos. Comprimido a ~12 págs total del docx."""
    cal, rms = load_cal_and_rms()
    r_cal = cal.get("r_cal", 0.005037)
    c_cal = cal.get("c_cal", 0.010355)

    # 7.1 Calibración + ecuación (4)
    add_body(
        doc,
        "La calibración se realizó en el corral C2 (K=100) con cinco "
        "réplicas para cada parámetro. La regeneración r se obtuvo "
        "ajustando la curva G(t) → K sin ovejas, y el consumo c "
        "ajustando la EDO (2) con N=5 manteniendo r fijo. Los valores "
        f"resultaron r = {r_cal:.6f} 1/s y c = {c_cal:.6f} "
        "bloques/(oveja·s), con una varianza inferior al 10 % entre "
        "réplicas. La predicción G*₊ = 88.4 coincide con el valor "
        "experimental G_final = 90.2 ± 2.3 dentro de ±1σ. La ecuación "
        "(4) presenta el modelo calibrado del proyecto."
    )
    add_equation_latex(
        doc,
        rf"\frac{{dG}}{{dt}} = {r_cal:.6f}\,G\left(1 - \frac{{G}}{{K}}\right) - {c_cal:.6f}\,N",
        4,
    )

    # Tabla 1 — calibración + N_crit por corral
    n_crit_C1 = r_cal * 25 / (4 * c_cal)
    n_crit_C2 = r_cal * 100 / (4 * c_cal)
    n_crit_C3 = r_cal * 225 / (4 * c_cal)
    add_table_from_rows(
        doc,
        ["Corral", "K", "N_crit", "Tipo"],
        [
            ["C1 (5×5)", "25", f"{n_crit_C1:.2f}", "Regular"],
            ["C2 (10×10)", "100", f"{n_crit_C2:.2f}", "Regular (referencia)"],
            ["C3 (15×15)", "225", f"{n_crit_C3:.2f}", "Regular"],
            ["C4 (12×12)", "100", f"{n_crit_C2:.2f}", "Obstáculo: agua"],
            ["C5 (12×12)", "100", f"{n_crit_C2:.2f}", "Obstáculo: hojas"],
            ["C6 (12×12)", "100", f"{n_crit_C2:.2f}", "Obstáculo: vallas"],
        ],
        "Tabla 1. Capacidad K y N crítico teórico por corral.",
    )

    # 7.2 Subcrítico OK + Figura 1
    add_body(
        doc,
        "En los escenarios subcríticos (esc 1, 4, 5, 7, 10 y 11), donde "
        "N < N_crit, la EDO (2) ajustó adecuadamente: G descendió y se "
        "estabilizó cerca de G*₊ con RMS entre 1.5 y 8.6 bloques. La "
        "Figura 1 muestra trayectorias normalizadas en C1/C2/C3 — al "
        "preservar N/K, los tres corrales convergen a la misma fracción."
    )
    add_picture_with_caption(
        doc,
        GRAFICAS_DIR / "grupo_A_K_creciente.png",
        "Figura 1. Efecto del tamaño K en régimen subcrítico (C1, C2, C3) "
        "con eje normalizado G/K.",
    )

    # 7.3 Defecto supercrítico + 7.4 refinamiento Holling
    add_body(
        doc,
        "Los escenarios supercríticos (esc 3, 6 y 9) revelaron un defecto "
        "de la EDO original: el modelo predice colapso a G = 0 cuando "
        "cN > rK/4, pero los datos mostraron mesetas estables en G ≈ 13, "
        "56 y 158 respectivamente. La explicación física es que cuando el "
        "pasto se vuelve escaso, las ovejas pierden tiempo buscando "
        "bloques aislados y el consumo efectivo cae, comportamiento que "
        "el término constante −cN no captura. Para corregirlo se introdujo "
        "la respuesta funcional Tipo II de Holling (Holling, 1959), "
        "ecuación (5), con una constante h de semi-saturación; cuando "
        "G ≫ h se recupera la ecuación (2), y cuando G → 0 el consumo "
        "se frena suavemente."
    )
    add_equation_latex(
        doc,
        r"\frac{dG}{dt} = rG\left(1 - \frac{G}{K}\right) - \frac{cNG}{G + h}",
        5,
    )
    add_body(
        doc,
        "El equilibrio del modelo refinado satisface la cuadrática (6) "
        "tras descartar la raíz trivial G = 0."
    )
    add_equation_latex(
        doc,
        r"G^2 - (K - h)G + K\left(\frac{cN}{r} - h\right) = 0",
        6,
    )
    add_body(
        doc,
        "La literatura ecológica reporta h/K entre 0.1 y 0.5 según la "
        "dificultad de búsqueda (Real, 1977; Turchin, 2003); aquí se "
        "adoptó h = K/2, valor que ajusta consistentemente las tres "
        "mesetas supercríticas y refleja la naturaleza discreta del "
        "corral (bloques individuales, no recurso continuo). La Figura 2 "
        "compara ambos modelos en los escenarios supercríticos."
    )
    add_picture_with_caption(
        doc,
        GRAFICAS_DIR / "comparacion_modelos.png",
        "Figura 2. Comparación de modelos en régimen supercrítico (esc 3, "
        "6, 9). La EDO original colapsa a cero; Holling reproduce las "
        "mesetas observadas.",
        width_inches=6.5,
    )

    # Tabla 2 — RMS comparativo
    if not rms.empty:
        rms_rows = []
        for _, row in rms.iterrows():
            esc = int(row["escenario"])
            corral = row["label"]
            N = int(row["N"])
            reg = row["regimen"]
            rms_o = row["RMS_orig"]
            rms_h = row["RMS_holling"]
            mejora = row.get("mejora_pct", float("nan"))
            try:
                mejora_str = f"{mejora:+.1f} %"
            except Exception:
                mejora_str = "-"
            rms_rows.append([
                str(esc), corral, str(N), reg,
                f"{rms_o:.2f}", f"{rms_h:.2f}", mejora_str,
            ])
        add_table_from_rows(
            doc,
            ["Esc", "Corral", "N", "Régimen", "RMS orig.", "RMS Holling",
             "Mejora"],
            rms_rows,
            "Tabla 2. RMS por escenario, EDO original vs Holling Tipo II.",
        )

    # 7.5 Obstáculos + datos extra C4/C6 N variado
    add_body(
        doc,
        "Los escenarios 10, 11 y 12 (corrales C4 agua, C5 hojas y C6 "
        "vallas, todos K=100 y N=10) aíslan el efecto del obstáculo, "
        "manteniendo el escenario 5 como control. Los 44 bloques de "
        "obstáculo se diseñaron para preservar el área efectiva de "
        "césped equivalente a C2 10×10. La Figura 3 superpone las cuatro "
        "trayectorias. C4 (agua) se comporta como C2: las ovejas nadan "
        "sin perder tiempo significativo. C5 (hojas) muestra equilibrio "
        "ligeramente inferior, posiblemente por el sombreado que afecta "
        "el nivel de luz mínimo para la regeneración (Minecraft Wiki, "
        "2024). C6 (vallas) presenta la desviación más fuerte: G_final ≈ "
        "54 frente al G*₊ predicho de 71 bloques. Las vallas no son "
        "saltables y las ovejas deben rodearlas, dejando rincones "
        "geométricamente inaccesibles. La capacidad de regeneración se "
        "conserva pero la capacidad efectiva de consumo es menor, "
        "asimetría que ningún modelo simétrico captura."
    )
    add_picture_with_caption(
        doc,
        GRAFICAS_DIR / "grupo_B_obstaculos.png",
        "Figura 3. Efecto del obstáculo (esc 5, 10, 11 y 12) con K=100 "
        "y N=10.",
    )
    add_body(
        doc,
        "Como evidencia adicional, en corridas exploratorias con N muy "
        "alto se observó que C4 (N = 82–90) mantiene mesetas pequeñas "
        "(G_final ≈ 6–23), consistentes con la saturación de Holling, "
        "mientras que C6 (N = 75–80) colapsa casi por completo "
        "(G_final = 1), confirmando que en C6 la limitante no es la "
        "respuesta funcional sino la geometría: bajo presión alta de "
        "consumo, las zonas inaccesibles dejan de regenerarse efectivamente."
    )

    # 7.6 Bifurcación + Figura 4
    add_body(
        doc,
        "La Figura 4 visualiza la bifurcación silla-nodo en C3 (K=225) "
        "barriendo los tres regímenes — subcrítico (N=10, G ≈ 200), "
        "cerca-crítico (N=25 próximo a N_crit ≈ 27) y supercrítico "
        "(N=40), donde el modelo Holling predice una meseta baja en "
        "lugar del colapso de la EDO original."
    )
    add_picture_with_caption(
        doc,
        GRAFICAS_DIR / "grupo_C_bifurcacion.png",
        "Figura 4. Bifurcación silla-nodo en C3 (K=225) con N = 10, 25, 40.",
    )


def _render_mejoras(doc: Document) -> None:
    add_body(
        doc,
        "Si bien el modelo refinado con respuesta funcional Tipo II "
        "reproduce la mayoría de escenarios, varios aspectos quedan "
        "abiertos como trabajo futuro. Primero, calibrar h "
        "explícitamente mediante un experimento dedicado (variando N "
        "finamente y ajustando c y h en conjunto) en lugar de asumirlo "
        "por literatura, y verificar si h escala con K. Segundo, extender "
        "el modelo a una forma asimétrica que distinga K de regeneración "
        "de K efectiva de consumo, necesaria para corrales con "
        "obstáculos no atravesables como C6. Tercero, complementar la "
        "EDO global con una simulación agente-basada del movimiento de "
        "las ovejas para reproducir efectos espaciales finos en corrales "
        "complejos. Cuarto, estudiar la sensibilidad de r al parámetro "
        "random_tick_speed del motor y validar la robustez de la "
        "calibración entre versiones del juego. Finalmente, automatizar "
        "la construcción de los corrales mediante un datapack facilitaría "
        "la reproducibilidad y la extensión a más configuraciones."
    )


def _render_conclusiones(doc: Document) -> None:
    add_body(
        doc,
        "Ambos objetivos específicos se cumplieron. La calibración "
        "controlada en C2 produjo r = 0.005037 1/s y c = 0.010355 "
        "bloques/(oveja·s) con varianza inferior al 10 %; la predicción "
        "G*₊ = 88.4 coincide con el valor experimental 90.2 ± 2.3, "
        "validando la consistencia interna del modelo. Los 12 escenarios "
        "cubrieron los tres regímenes esperados, identificaron "
        "empíricamente la bifurcación silla-nodo y cuantificaron el "
        "efecto del obstáculo: agua y hojas afectan marginalmente la "
        "dinámica, mientras que las vallas reducen la capacidad efectiva "
        "de consumo por zonas geométricamente inaccesibles."
    )
    add_body(
        doc,
        "La EDO logística con extracción constante (ecuación 2) describe "
        "adecuadamente el régimen subcrítico (RMS < 9 bloques) pero "
        "predice un colapso falso a G = 0 en régimen supercrítico; los "
        "datos motivaron la respuesta funcional Tipo II de Holling "
        "(ecuación 5), que mejora el RMS hasta 89 % en los escenarios "
        "problemáticos al capturar la saturación del consumo cuando el "
        "pasto es escaso. La asunción de r y c constantes entre corrales "
        "se sostuvo en regulares y en obstáculos atravesables, pero no en "
        "C6 (vallas), donde se requiere un modelo asimétrico. En "
        "perspectiva metodológica, este trabajo muestra que incluso en un "
        "sistema con reglas algorítmicas conocidas, un modelo clásico "
        "puede presentar defectos que sólo se evidencian al contrastarlo "
        "con datos, y que la iteración teoría–experimento es lo que "
        "permite construir modelos verdaderamente útiles."
    )


def build_pf(out: Path) -> None:
    doc = Document()
    setup_styles(doc)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.line_spacing = 2.0
    run = title_p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(14)

    authors_p = doc.add_paragraph()
    authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_p.paragraph_format.line_spacing = 2.0
    authors_p.add_run(AUTHORS).italic = True

    # 1. Introducción
    add_heading(doc, "Introducción", level=1)
    add_body(doc, INTRO)

    # 2. Estado del arte
    add_heading(doc, "Estado del arte", level=1)
    for block in ESTADO_BLOCKS:
        if block[0] == "text":
            add_body(doc, block[1])
        elif block[0] == "eq":
            add_equation_latex(doc, block[1], block[2])

    # 3. Planteamiento de la propuesta
    add_heading(doc, "Planteamiento de la propuesta", level=1)
    add_body(doc, PROPUESTA)

    # 4. Objetivos
    add_heading(doc, "Objetivos", level=1)
    add_heading(doc, "Objetivo general", level=2)
    add_body(doc, OBJETIVO_GENERAL)
    add_heading(doc, "Objetivos específicos", level=2)
    for i, obj in enumerate(OBJETIVOS_ESPEC, 1):
        p = doc.add_paragraph(f"{i}. {obj}")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 2.0

    # 5. Metodología
    add_heading(doc, "Metodología", level=1)
    step_idx = 0
    for block in METODOLOGIA_BLOCKS:
        if block[0] == "step":
            step_idx += 1
            p = doc.add_paragraph(f"{step_idx}. {block[1]}")
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 2.0
        elif block[0] == "eq":
            add_equation_latex(doc, block[1], block[2])

    # 6. Justificación
    add_heading(doc, "Justificación", level=1)
    for p in JUSTIFICACION_PARRAFOS:
        add_body(doc, p)

    # 7. Resultados
    add_heading(doc, "Resultados", level=1)
    _render_resultados(doc)

    # 8. Posibles mejoras y trabajo futuro
    add_heading(doc, "Posibles mejoras y trabajo futuro", level=1)
    _render_mejoras(doc)

    # 9. Conclusiones
    add_heading(doc, "Conclusiones", level=1)
    _render_conclusiones(doc)

    # 10. Referencias
    add_heading(doc, "Referencias", level=1)
    for runs in REFERENCIAS:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 2.0
        pf = p.paragraph_format
        pf.first_line_indent = Pt(-18)
        pf.left_indent = Pt(18)
        for text, italic in runs:
            r = p.add_run(text)
            r.italic = italic

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Saved {out}")


def build_bepf(out: Path) -> None:
    doc = Document()
    setup_styles(doc)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.line_spacing = 2.0
    run = title_p.add_run(
        "BITÁCORA-EXPOSICIÓN PROYECTO FINAL — VIDEO"
    )
    run.bold = True
    run.font.size = Pt(14)

    authors_p = doc.add_paragraph()
    authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_p.paragraph_format.line_spacing = 2.0
    authors_p.add_run(AUTHORS).italic = True

    spacer = doc.add_paragraph("")
    spacer.paragraph_format.line_spacing = 2.0

    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.line_spacing = 2.0
    label.add_run("Enlace al video:").bold = True

    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link_p.paragraph_format.line_spacing = 2.0
    run = link_p.add_run("[REEMPLAZAR_CON_URL_DEL_VIDEO]")
    run.font.size = Pt(12)
    run.italic = True

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Saved {out}")


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    templates = root / "templates"
    build_pf(
        templates
        / "PF-G2-Bello Ballen, Cristancho Niño, Gordillo Meneses, Vargas Clavijo.docx"
    )
    build_bepf(
        templates
        / "BEPF-G2-Bello Ballen, Cristancho Niño, Gordillo Meneses, Vargas Clavijo.docx"
    )


if __name__ == "__main__":
    main()
